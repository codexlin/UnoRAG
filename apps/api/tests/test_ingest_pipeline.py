"""Phase A–C ingest pipeline unit tests."""

from __future__ import annotations

from pathlib import Path

import pytest

from app.services.ingest.chunker import ChunkerConfig, chunk_document
from app.services.ingest.ir import NodeType, SplitStrategy, format_page_label
from app.services.ingest.parsers.md import parse_markdown
from app.services.ingest.parsers.pdf import classify_page
from app.services.ingest.parsers.txt import parse_txt
from app.services.ingest.pipeline import prepare_ingest
from app.settings import Settings

FIXTURES = Path(__file__).parent / "fixtures"


def test_handbook_md_no_cross_h2_and_preamble() -> None:
	content = (FIXTURES / "handbook.md").read_bytes()
	doc = parse_markdown(
		content=content,
		filename="handbook.md",
		title="员工手册",
	)
	assert any(n.type == NodeType.HEADING and "第3章" in (n.text or "") for n in doc.nodes)
	assert any(n.type == NodeType.TABLE and n.table_id for n in doc.nodes)

	chunks = chunk_document(doc, config=ChunkerConfig(chunk_size=500, chunk_overlap=40))
	assert chunks
	# 每个带 section 的正文块应有 preamble
	body_chunks = [c for c in chunks if c.body and c.split_strategy != SplitStrategy.TABLE]
	assert body_chunks
	assert all(c.preamble for c in body_chunks)

	# 不跨 H2：含「第12条 病假」的块不应同时含「第4章」
	leave_chunks = [c for c in chunks if "病假" in c.body]
	assert leave_chunks
	for chunk in leave_chunks:
		assert "第4章" not in chunk.body
		assert chunk.section_path is None or "第3章" in (chunk.section_path or "")


def test_sibling_h3_section_paths_do_not_cross() -> None:
	"""同级 ### 必须各自成段，避免正文挂到下一节 path（CMSTOP → 璞华串号）。"""
	content = """# 简历

## 工作经历

### CMSTOP
Monorepo 迁移与规范。

### 璞华国际
LangChain 企业助手。
""".encode("utf-8")
	doc = parse_markdown(content=content, filename="resume.md", title="简历")
	chunks = chunk_document(doc, config=ChunkerConfig(chunk_size=500, chunk_overlap=40))
	cmstop = [c for c in chunks if "Monorepo" in c.body]
	puhua = [c for c in chunks if "LangChain" in c.body]
	assert cmstop
	assert puhua
	assert all(c.section_path and "CMSTOP" in c.section_path for c in cmstop)
	assert all("璞华" not in (c.section_path or "") for c in cmstop)
	assert all(c.section_path and "璞华" in c.section_path for c in puhua)
	assert all("Monorepo" not in c.body for c in puhua)


def test_txt_parser_paragraphs() -> None:
	content = "第一段内容。\n\n第二段内容。".encode("utf-8")
	doc = parse_txt(content=content, filename="a.txt", title="a")
	assert len(doc.nodes) >= 2


def test_prepare_ingest_v2_md(monkeypatch: pytest.MonkeyPatch) -> None:
	settings = Settings(
		ingest_pipeline="v2",
		ask_mode="stub",
		metadata_backend="json",
	)
	content = (FIXTURES / "handbook.md").read_bytes()
	prepared = prepare_ingest(
		settings=settings,
		filename="handbook.md",
		content=content,
		library_id="lib-hr",
	)
	assert prepared.pipeline == "v2"
	assert prepared.chunks
	assert any(c.section_path and "考勤" in (c.section_path or "") for c in prepared.chunks)


def test_prepare_ingest_legacy_flag() -> None:
	settings = Settings(ingest_pipeline="legacy", ask_mode="stub", metadata_backend="json")
	prepared = prepare_ingest(
		settings=settings,
		filename="note.md",
		content=b"# Hi\n\nHello world",
		library_id="lib-hr",
	)
	assert prepared.pipeline == "legacy"
	assert all(c.split_strategy == SplitStrategy.CHAR_WINDOW for c in prepared.chunks)


def test_pdf_page_classify() -> None:
	assert classify_page(char_count=500, image_area_ratio=0.0, image_count=0) == "text"
	assert classify_page(char_count=5, image_area_ratio=0.5, image_count=1) == "suspect_scan"
	assert classify_page(char_count=80, image_area_ratio=0.5, image_count=2) == "complex"


def test_prepare_ingest_applies_ocr_override(monkeypatch: pytest.MonkeyPatch) -> None:
	"""scan_handling force_ocr/disabled must reach PDF parse options."""
	from app.services.ingest import router as ingest_router
	from app.services.ingest.ir import DocumentIR, ParserReport, Node, NodeType

	seen: dict[str, object] = {}

	def fake_get_ocr_adapter(*, enabled: bool):
		seen["adapter_enabled"] = enabled
		return None

	def fake_get_vlm_adapter(*, enabled: bool, **_kwargs):
		seen["vlm_adapter_enabled"] = enabled
		return None

	def fake_parse_pdf_routed(**kwargs):
		opts = kwargs["options"]
		seen["options_ocr"] = opts.ocr_enabled
		seen["options_vlm"] = opts.vlm_enabled
		seen["enhanced_parser_allowed"] = kwargs["enhanced_parser_allowed"]
		return DocumentIR(
			id="doc-ocr",
			title="scan",
			filename="scan.pdf",
			source_format="pdf",
			nodes=[
				Node(
					id="n1",
					type=NodeType.PAGE,
					path="第1页",
					page_start=1,
					page_end=1,
					text="hello",
				)
			],
			parser_report=ParserReport(source_format="pdf", parser="test"),
		)

	monkeypatch.setattr(ingest_router, "get_ocr_adapter", fake_get_ocr_adapter)
	monkeypatch.setattr(ingest_router, "get_vlm_adapter", fake_get_vlm_adapter)
	monkeypatch.setattr(ingest_router, "parse_pdf_routed", fake_parse_pdf_routed)

	settings = Settings(
		ingest_pipeline="v2",
		ask_mode="stub",
		metadata_backend="json",
		ocr_enabled=False,
		vlm_enabled=True,
	)
	prepared = prepare_ingest(
		settings=settings,
		filename="scan.pdf",
		content=b"%PDF-1.4 fake",
		library_id="lib-ocr",
		ocr_enabled=True,
		enhanced_parser_allowed=True,
	)
	assert prepared.pipeline == "v2"
	assert seen["adapter_enabled"] is True
	assert seen["options_ocr"] is True
	assert seen["options_vlm"] is True
	assert seen["enhanced_parser_allowed"] is True

	seen.clear()
	prepare_ingest(
		settings=Settings(
			ingest_pipeline="v2",
			ask_mode="stub",
			metadata_backend="json",
			ocr_enabled=True,
			vlm_enabled=True,
		),
		filename="scan.pdf",
		content=b"%PDF-1.4 fake",
		library_id="lib-ocr",
		ocr_enabled=False,
		enhanced_parser_allowed=False,
	)
	assert seen["adapter_enabled"] is False
	assert seen["vlm_adapter_enabled"] is False
	assert seen["options_ocr"] is False
	assert seen["options_vlm"] is False
	assert seen["enhanced_parser_allowed"] is False


def test_format_page_label_range() -> None:
	assert format_page_label(1) == "p.1"
	assert format_page_label(2, 4) == "p.2-4"


def test_docx_parser_optional() -> None:
	pytest.importorskip("docx")
	from docx import Document
	import io

	from app.services.ingest.parsers.docx import parse_docx

	buf = io.BytesIO()
	d = Document()
	d.add_heading("制度", level=1)
	d.add_heading("请假", level=2)
	d.add_paragraph("病假须提交证明。")
	d.save(buf)
	doc = parse_docx(content=buf.getvalue(), filename="policy.docx", title="制度")
	assert any(n.type == NodeType.HEADING for n in doc.nodes)
	chunks = chunk_document(doc)
	assert chunks
	assert any("病假" in c.body for c in chunks)
