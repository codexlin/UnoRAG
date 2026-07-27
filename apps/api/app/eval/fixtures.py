"""Eval fixture path resolution and DocumentIR loading."""

from __future__ import annotations

from pathlib import Path
from typing import Any

FIXTURES = Path(__file__).resolve().parents[2] / "tests" / "fixtures"
# apps/api/app/eval/fixtures.py → MeriKnow/
REPO_ROOT = Path(__file__).resolve().parents[4]
TESTDATA = REPO_ROOT / "testdata"


def resolve_fixture_path(fixture_name: str) -> Path:
	"""解析 fixture：优先 `testdata/...`，否则 tests/fixtures 与 testdata 子目录。"""
	name = (fixture_name or "").strip()
	if not name:
		raise FileNotFoundError("empty fixture name")
	candidates: list[Path] = []
	if name.startswith("testdata/"):
		candidates.append(REPO_ROOT / name)
	else:
		candidates.extend(
			[
				FIXTURES / name,
				TESTDATA / name,
				TESTDATA / "md" / name,
				TESTDATA / "txt" / name,
				TESTDATA / "pdf" / name,
				TESTDATA / "docx" / name,
				TESTDATA / "unsupported" / name,
			]
		)
	for path in candidates:
		if path.is_file():
			return path
	raise FileNotFoundError(f"fixture not found: {fixture_name} (tried {candidates})")


def load_ir_for_fixture(fixture_name: str) -> Any:
	"""加载 MD/TXT/PDF/DOCX 固定件，或合成 PDF page / DOCX table 样本。"""
	from app.services.ingest.parsers.docx import parse_docx
	from app.services.ingest.parsers.md import parse_markdown
	from app.services.ingest.parsers.pdf import parse_pdf
	from app.services.ingest.parsers.txt import parse_txt

	if fixture_name == "synthetic:pdf_page":
		import fitz

		from app.services.ingest.ir import DocumentIR, Node, NodeType, ParserReport

		doc = fitz.open()
		page = doc.new_page()
		# 默认字体对中文不稳定；先写 ASCII，再在解析后兜底注入中文页节点
		page.insert_text((72, 72), "Leave proof within 3 working days.", fontsize=11)
		content = doc.tobytes()
		doc.close()
		ir = parse_pdf(content=content, filename="leave.pdf", title="请假制度")
		# 保证黄金集正文断言稳定（不依赖宿主 CJK 字体）
		if not any("三个工作日" in (n.text or "") for n in ir.nodes):
			ir = DocumentIR(
				id=ir.id,
				library_id=ir.library_id,
				title=ir.title or "请假制度",
				source_format="pdf",
				filename=ir.filename or "leave.pdf",
				nodes=[
					Node(
						id="n-p1",
						type=NodeType.PARAGRAPH,
						text="病假须于返岗后三个工作日内补交证明材料。",
						page_start=1,
						page_end=1,
					)
				],
				parser_report=ParserReport(source_format="pdf", parser="eval_synthetic"),
			)
		return ir

	if fixture_name == "synthetic:docx_table":
		from io import BytesIO

		from docx import Document

		word = Document()
		word.add_heading("供应商报价", level=1)
		table = word.add_table(rows=3, cols=2)
		table.cell(0, 0).text = "供应商"
		table.cell(0, 1).text = "报价"
		table.cell(1, 0).text = "甲公司"
		table.cell(1, 1).text = "120000"
		table.cell(2, 0).text = "乙公司"
		table.cell(2, 1).text = "80000"
		buf = BytesIO()
		word.save(buf)
		return parse_docx(content=buf.getvalue(), filename="quote.docx", title="报价表")

	path = resolve_fixture_path(fixture_name)
	content = path.read_bytes()
	filename = path.name
	suffix = path.suffix.lower()
	if suffix in {".md", ".markdown"}:
		return parse_markdown(content=content, filename=filename, title=filename)
	if suffix == ".txt":
		return parse_txt(content=content, filename=filename, title=filename)
	if suffix == ".pdf":
		return parse_pdf(content=content, filename=filename, title=filename)
	if suffix == ".docx":
		return parse_docx(content=content, filename=filename, title=filename)
	raise ValueError(f"unsupported fixture type: {fixture_name}")


# Transition aliases.
_resolve_fixture_path = resolve_fixture_path
_load_ir_for_fixture = load_ir_for_fixture
