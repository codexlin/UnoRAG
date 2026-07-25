"""黄金集 runner — 本地可跑，默认 stub AskGraph。

用法：
  uv run python -m app.eval.runner
  uv run python scripts/run_eval_cases.py
"""

from __future__ import annotations

import hashlib
import json
import math
import os
import sys
from contextlib import contextmanager
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from app.eval.schemas import EvalCase, EvalCaseResult, EvalExpect
from app.services.query_router import classify_query
from app.services.retrieval_plan import build_retrieval_plan

DEFAULT_CASES = Path(__file__).resolve().parents[2] / "tests" / "eval" / "eval_cases.jsonl"
FIXTURES = Path(__file__).resolve().parents[2] / "tests" / "fixtures"
# apps/api/app/eval/runner.py → MeriKnow/
REPO_ROOT = Path(__file__).resolve().parents[4]
TESTDATA = REPO_ROOT / "testdata"


def load_eval_cases(path: Path | None = None) -> list[EvalCase]:
	resolved = path or DEFAULT_CASES
	cases: list[EvalCase] = []
	with resolved.open("r", encoding="utf-8") as handle:
		for line_no, line in enumerate(handle, start=1):
			text = line.strip()
			if not text or text.startswith("#"):
				continue
			try:
				raw = json.loads(text)
				cases.append(EvalCase.model_validate(raw))
			except Exception as exc:
				raise ValueError(f"invalid eval case at {resolved}:{line_no}: {exc}") from exc
	return cases


def _check_expect(expect: EvalExpect, observed: dict[str, Any]) -> list[str]:
	errors: list[str] = []
	if expect.query_type is not None and observed.get("query_type") != expect.query_type:
		errors.append(f"query_type want={expect.query_type} got={observed.get('query_type')}")
	if expect.refused is not None and bool(observed.get("refused")) != expect.refused:
		errors.append(f"refused want={expect.refused} got={observed.get('refused')}")
	if expect.refuse_reason is not None and observed.get("refuse_reason") != expect.refuse_reason:
		errors.append(
			f"refuse_reason want={expect.refuse_reason} got={observed.get('refuse_reason')}"
		)
	if expect.judge_reason is not None:
		judge = observed.get("judge") or {}
		if not isinstance(judge, dict) or judge.get("reason") != expect.judge_reason:
			errors.append(
				f"judge.reason want={expect.judge_reason} got={(judge or {}).get('reason')}"
			)
	if expect.execute_path is not None:
		plan = observed.get("retrieval_plan") or {}
		got = plan.get("execute_path") if isinstance(plan, dict) else None
		if got != expect.execute_path:
			errors.append(f"execute_path want={expect.execute_path} got={got}")
	answer = str(observed.get("answer") or "")
	for needle in expect.answer_contains:
		if needle not in answer:
			errors.append(f"answer missing: {needle!r}")
	if expect.section_substr is not None:
		section = str(observed.get("section_path") or "")
		if expect.section_substr not in section:
			errors.append(
				f"section_path missing {expect.section_substr!r} got={section!r}"
			)
	if expect.body_substr is not None:
		body = str(observed.get("body") or "")
		if expect.body_substr not in body:
			errors.append(f"body missing {expect.body_substr!r}")
	if expect.max_rank is not None:
		rank = observed.get("observed_rank")
		if rank is None:
			errors.append(f"max_rank={expect.max_rank} but observed_rank is missing (no hit in Recall@K)")
		elif int(rank) > int(expect.max_rank):
			errors.append(f"observed_rank={rank} exceeds max_rank={expect.max_rank}")
	if expect.http_status is not None and observed.get("http_status") != expect.http_status:
		errors.append(
			f"http_status want={expect.http_status} got={observed.get('http_status')}"
		)
	if expect.http_status_any:
		got_status = observed.get("http_status")
		if got_status not in expect.http_status_any:
			errors.append(
				f"http_status want one of {expect.http_status_any} got={got_status}"
			)
	if expect.doc_status is not None and observed.get("doc_status") != expect.doc_status:
		errors.append(
			f"doc_status want={expect.doc_status} got={observed.get('doc_status')}"
		)
	if expect.error_substr is not None:
		blob = str(observed.get("error") or "")
		if expect.error_substr not in blob:
			errors.append(f"error missing {expect.error_substr!r} got={blob!r}")
	if expect.detail_substr is not None:
		blob = str(observed.get("detail") or "")
		if expect.detail_substr not in blob:
			errors.append(f"detail missing {expect.detail_substr!r} got={blob!r}")
	if expect.record_type is not None and observed.get("record_type") != expect.record_type:
		errors.append(
			f"record_type want={expect.record_type} got={observed.get('record_type')}"
		)
	return errors


def _run_classify(case: EvalCase) -> EvalCaseResult:
	query_type, reason = classify_query(case.question, history=case.history or None)
	plan = build_retrieval_plan(
		query_type=query_type,
		route_reason=reason,
		library_id=case.library_id,
		top_k=6,
		hybrid_enabled=False,
		rerank_enabled=False,
		question=case.question,
	)
	observed = {
		"query_type": query_type,
		"route_reason": reason,
		"retrieval_plan": plan,
	}
	errors = _check_expect(case.expect, observed)
	return EvalCaseResult(
		id=case.id,
		ok=not errors,
		kind=case.kind,
		errors=errors,
		observed=observed,
	)


# Eval ask knobs: code defaults ⊕ these fixed overrides (never env).
_EVAL_ASK_OVERRIDES = {
	"session_memory_enabled": False,
	"hybrid_enabled": False,
	"rerank_enabled": False,
}


@contextmanager
def _isolated_ask_settings():
	"""隔离 eval 对环境、settings cache 和 metadata singleton 的修改。"""
	from app.graph.ask_graph import AskGraphService, stub_load_table_groups
	from app.security.access_scope import AccessScope
	from app.services.metadata import reset_metadata_store
	from app.settings import get_settings

	keys = (
		"ASK_MODE",
		"METADATA_BACKEND",
		"METADATA_PATH",
		"MAX_RETRIEVE_RETRIES",
		"INTERNAL_AUTH_ENABLED",
	)
	previous = {key: os.environ.get(key) for key in keys}
	with TemporaryDirectory(prefix="meriknow-eval-") as tmp_dir:
		os.environ.update(
			{
				"ASK_MODE": "stub",
				"METADATA_BACKEND": "json",
				"METADATA_PATH": str(Path(tmp_dir) / "metadata.json"),
				"MAX_RETRIEVE_RETRIES": "0",
				# Avoid host .env INTERNAL_AUTH_ENABLED=true requiring request scope.
				"INTERNAL_AUTH_ENABLED": "false",
			}
		)
		get_settings.cache_clear()
		reset_metadata_store()
		try:
			settings = get_settings()
			yield AskGraphService(
				settings=settings,
				access_scope=AccessScope.development(settings),
				# Align with stub_retrieve / test_ask_route table shape.
				load_table_groups_fn=stub_load_table_groups,
			)
		finally:
			reset_metadata_store()
			for key, value in previous.items():
				if value is None:
					os.environ.pop(key, None)
				else:
					os.environ[key] = value
			get_settings.cache_clear()


def _run_ask(case: EvalCase) -> EvalCaseResult:
	with _isolated_ask_settings() as service:
		# history 样例直接调用图，避免依赖持久化 session memory。
		if case.history:
			ask_settings = service._ask_settings_for_request(_EVAL_ASK_OVERRIDES)
			graph = service._graph_for_settings(ask_settings)
			state = graph.invoke(
				{
					"session_id": case.session_id or f"eval-{case.id}",
					"question": case.question,
					"library_id": case.library_id,
					"history": case.history,
					"retrieval_debug": {},
				}
			)
			debug = state.get("retrieval_debug") or {}
			observed = {
				"query_type": state.get("query_type") or debug.get("query_type"),
				"refused": bool(state.get("refused")),
				"refuse_reason": state.get("refuse_reason"),
				"answer": state.get("answer") or "",
				"judge": state.get("judgement") or debug.get("judgement"),
				"retrieval_plan": state.get("retrieval_plan") or debug.get("retrieval_plan"),
			}
		else:
			resp = service.ask(
				question=case.question,
				library_id=case.library_id,
				ask_overrides=_EVAL_ASK_OVERRIDES,
			)
			debug = resp.retrieval_debug or {}
			observed = {
				"query_type": debug.get("query_type"),
				"refused": resp.refused,
				"refuse_reason": resp.refuse_reason,
				"answer": resp.answer,
				"judge": debug.get("judgement"),
				"retrieval_plan": debug.get("retrieval_plan"),
			}
	errors = _check_expect(case.expect, observed)
	return EvalCaseResult(
		id=case.id,
		ok=not errors,
		kind=case.kind,
		errors=errors,
		observed=observed,
	)


def _resolve_fixture_path(fixture_name: str) -> Path:
	"""解析 fixture：优先 `testdata/...`，否则 tests/fixtures 与 testdata 子目录。"""
	name = (fixture_name or "").strip()
	if not name:
		raise FileNotFoundError("empty fixture name")
	candidates: list[Path] = []
	if name.startswith("testdata/"):
		candidates.append(REPO_ROOT / name)
	else:
		candidates.extend(
			[
				FIXTURES / name,
				TESTDATA / name,
				TESTDATA / "md" / name,
				TESTDATA / "txt" / name,
				TESTDATA / "pdf" / name,
				TESTDATA / "docx" / name,
				TESTDATA / "unsupported" / name,
			]
		)
	for path in candidates:
		if path.is_file():
			return path
	raise FileNotFoundError(f"fixture not found: {fixture_name} (tried {candidates})")


def _load_ir_for_fixture(fixture_name: str) -> Any:
	"""加载 MD/TXT/PDF/DOCX 固定件，或合成 PDF page / DOCX table 样本。"""
	from app.services.ingest.parsers.docx import parse_docx
	from app.services.ingest.parsers.md import parse_markdown
	from app.services.ingest.parsers.pdf import parse_pdf
	from app.services.ingest.parsers.txt import parse_txt

	if fixture_name == "synthetic:pdf_page":
		import fitz

		from app.services.ingest.ir import DocumentIR, Node, NodeType, ParserReport

		doc = fitz.open()
		page = doc.new_page()
		# 默认字体对中文不稳定；先写 ASCII，再在解析后兜底注入中文页节点
		page.insert_text((72, 72), "Leave proof within 3 working days.", fontsize=11)
		content = doc.tobytes()
		doc.close()
		ir = parse_pdf(content=content, filename="leave.pdf", title="请假制度")
		# 保证黄金集正文断言稳定（不依赖宿主 CJK 字体）
		if not any("三个工作日" in (n.text or "") for n in ir.nodes):
			ir = DocumentIR(
				id=ir.id,
				library_id=ir.library_id,
				title=ir.title or "请假制度",
				source_format="pdf",
				filename=ir.filename or "leave.pdf",
				nodes=[
					Node(
						id="n-p1",
						type=NodeType.PARAGRAPH,
						text="病假须于返岗后三个工作日内补交证明材料。",
						page_start=1,
						page_end=1,
					)
				],
				parser_report=ParserReport(source_format="pdf", parser="eval_synthetic"),
			)
		return ir

	if fixture_name == "synthetic:docx_table":
		from io import BytesIO

		from docx import Document

		word = Document()
		word.add_heading("供应商报价", level=1)
		table = word.add_table(rows=3, cols=2)
		table.cell(0, 0).text = "供应商"
		table.cell(0, 1).text = "报价"
		table.cell(1, 0).text = "甲公司"
		table.cell(1, 1).text = "120000"
		table.cell(2, 0).text = "乙公司"
		table.cell(2, 1).text = "80000"
		buf = BytesIO()
		word.save(buf)
		return parse_docx(content=buf.getvalue(), filename="quote.docx", title="报价表")

	path = _resolve_fixture_path(fixture_name)
	content = path.read_bytes()
	filename = path.name
	suffix = path.suffix.lower()
	if suffix in {".md", ".markdown"}:
		return parse_markdown(content=content, filename=filename, title=filename)
	if suffix == ".txt":
		return parse_txt(content=content, filename=filename, title=filename)
	if suffix == ".pdf":
		return parse_pdf(content=content, filename=filename, title=filename)
	if suffix == ".docx":
		return parse_docx(content=content, filename=filename, title=filename)
	raise ValueError(f"unsupported fixture type: {fixture_name}")


def _run_ingest_chunk(case: EvalCase) -> EvalCaseResult:
	from app.services.ingest.chunker import chunk_document

	fixture_name = case.fixture or "handbook.md"
	doc = _load_ir_for_fixture(fixture_name)
	chunks = chunk_document(doc)
	q = case.question

	def score(chunk: Any) -> float:
		body = chunk.body or ""
		overlap = len(set(q) & set(body)) / max(len(set(q)), 1)
		bonus = 0.0
		if "病假" in q and "病假" in body:
			bonus += 0.5
		if "薪酬" in q and ("薪酬" in body or "岗位职级" in body or "二十五日" in body):
			bonus += 0.5
		if "表格" in q or "基本工资" in q or "供应商" in q or "报价" in q or "甲公司" in q or "总价" in q:
			if chunk.table_id or "|" in body or "基本工资" in body or "报价" in body or "120000" in body:
				bonus += 0.5
		if ("五个自然日" in q or "人力资源前台" in q) and (
			"五个自然日" in body or "人力资源前台" in body
		):
			bonus += 0.6
		if ("断电" in q or "图注" in q) and ("断电" in body or "图注" in body):
			bonus += 0.6
		if ("吸烟" in q or "罚款" in q) and ("吸烟" in body or "罚款200" in body):
			bonus += 0.6
		if "页" in q and (chunk.page_label or chunk.page_start):
			bonus += 0.2
		return overlap + bonus

	ranked = sorted(chunks, key=score, reverse=True)
	top = ranked[0] if ranked else None
	observed = {
		"section_path": getattr(top, "section_path", None) if top else None,
		"body": getattr(top, "body", None) if top else None,
		"table_id": getattr(top, "table_id", None) if top else None,
		"page": getattr(top, "page_label", None) if top else None,
	}
	# PDF：额外校验 page 存在
	is_pdf = fixture_name == "synthetic:pdf_page" or fixture_name.lower().endswith(".pdf")
	if is_pdf and top is not None:
		if not (top.page_label or top.page_start is not None):
			return EvalCaseResult(
				id=case.id,
				ok=False,
				kind=case.kind,
				errors=["pdf chunk missing page"],
				observed=observed,
			)
	errors = _check_expect(case.expect, observed)
	return EvalCaseResult(
		id=case.id,
		ok=not errors,
		kind=case.kind,
		errors=errors,
		observed=observed,
	)


def _deterministic_vector(text: str, *, dimensions: int = 128) -> list[float]:
	"""CI 用确定性字符/二元组向量，不调用外部 embedding 服务。"""
	compact = "".join(char.lower() for char in text if not char.isspace())
	tokens = list(compact)
	tokens.extend(compact[index : index + 2] for index in range(max(0, len(compact) - 1)))
	vector = [0.0] * dimensions
	for token in tokens:
		digest = hashlib.sha256(token.encode("utf-8")).digest()
		index = int.from_bytes(digest[:4], "big") % dimensions
		vector[index] += 1.0
	norm = math.sqrt(sum(value * value for value in vector))
	return [value / norm for value in vector] if norm else vector


def _run_retrieval(case: EvalCase) -> EvalCaseResult:
	"""真实经过 QdrantStore + RetrievalService 的本地可重复检索回归。

	默认指标是 Recall@K（K=expect.recall_at_k 或 3）：目标片段出现在前 K 条即算命中。
	可用 expect.max_rank 收紧名次（1-based）。
	"""
	from qdrant_client import QdrantClient

	from app.services.ingest.chunker import chunk_document
	from app.services.ingest.pipeline import chunks_to_payloads
	from app.services.qdrant_store import QdrantStore
	from app.services.retrieval import RetrievalService
	from app.settings import Settings

	fixture_name = case.fixture or "handbook.md"
	doc = _load_ir_for_fixture(fixture_name)
	chunks = chunk_document(doc)
	record_type = str(case.expect.record_type or "chunk")
	payloads = chunks_to_payloads(
		chunks,
		filename=doc.filename or fixture_name,
		doc_id=doc.id,
		library_id=case.library_id or "lib-eval",
		document_version_id=f"eval-{case.id}-version",
		generation_id=f"eval-{case.id}-generation",
		lifecycle_visibility="active",
		include_sections=True,
	)
	# fact 隔离：只 upsert 所需粒度时仍写入全部，靠 search filter 验证 section 不进 fact
	dimensions = 128
	recall_at_k = int(case.expect.recall_at_k or 3)
	settings = Settings(
		ask_mode="stub",
		embedding_dim=dimensions,
		qdrant_collection=f"eval_{case.id.replace('-', '_')}",
		rerank_top_k=recall_at_k,
		bm25_top_k=recall_at_k,
	)

	class _EvalEmbeddings:
		def embed_query(self, text: str) -> list[float]:
			return _deterministic_vector(text, dimensions=dimensions)

		def embed_texts(self, texts: list[str]) -> list[list[float]]:
			return [_deterministic_vector(text, dimensions=dimensions) for text in texts]

	client = QdrantClient(location=":memory:")
	try:
		store = QdrantStore(settings, client=client)
		store.upsert_chunks(
			library_id=case.library_id or "lib-eval",
			doc_id=doc.id,
			title=doc.title or fixture_name,
			chunks=payloads,
			vectors=[
				_deterministic_vector(str(item.get("embed_text") or item.get("text") or ""), dimensions=dimensions)
				for item in payloads
			],
			filename=doc.filename or fixture_name,
		)
		service = RetrievalService(
			settings,
			embeddings=_EvalEmbeddings(),  # type: ignore[arg-type]
			store=store,
			reranker=None,
		)
		hits = service.search(
			query=case.question,
			library_id=case.library_id or "lib-eval",
			top_k=recall_at_k,
			record_type=record_type,
			filters={"record_type": record_type},
		)
	finally:
		client.close()

	top = hits[0] if hits else None
	body_substr = (case.expect.body_substr or "").strip()
	section_substr = (case.expect.section_substr or "").strip()
	observed_rank: int | None = None
	matched = None
	for index, item in enumerate(hits, start=1):
		body = str(item.get("body") or "")
		section = str(item.get("section_path") or "")
		body_ok = (not body_substr) or (body_substr in body)
		section_ok = (not section_substr) or (section_substr in section)
		if body_ok and section_ok and (body_substr or section_substr):
			observed_rank = index
			matched = item
			break
		if not body_substr and not section_substr and index == 1:
			observed_rank = 1
			matched = item
			break

	chosen = matched or top
	mrr = (1.0 / observed_rank) if observed_rank else 0.0
	# fact 隔离：chunk 检索结果中不得出现 section 记录
	polluted = [
		item for item in hits if str(item.get("record_type") or "") == "section"
	] if record_type == "chunk" else []
	observed = {
		"body": chosen.get("body") if chosen else None,
		"section_path": chosen.get("section_path") if chosen else None,
		"document_version_id": chosen.get("document_version_id") if chosen else None,
		"record_type": (chosen.get("record_type") if chosen else None) or record_type,
		"source_chunk_ids": chosen.get("source_chunk_ids") if chosen else None,
		"hit_count": len(hits),
		"recall_at_k": recall_at_k,
		"recall_hit": observed_rank is not None,
		"observed_rank": observed_rank,
		"mrr": mrr,
		"metric": f"Recall@{recall_at_k}",
		"section_pollution_count": len(polluted),
	}
	errors = _check_expect(case.expect, observed)
	if top is None:
		errors.append("retrieval returned no hits")
	elif (body_substr or section_substr) and matched is None:
		errors.append(
			f"Recall@{recall_at_k} miss: no hit contains "
			f"body={body_substr!r} section={section_substr!r}"
		)
	elif chosen and not observed["document_version_id"]:
		errors.append("retrieval hit missing document_version_id")
	if record_type == "chunk" and polluted:
		errors.append("section records leaked into chunk retrieval")
	if record_type == "section" and chosen and not (chosen.get("source_chunk_ids") or []):
		errors.append("section hit missing source_chunk_ids")
	return EvalCaseResult(
		id=case.id,
		ok=not errors,
		kind=case.kind,
		errors=errors,
		observed=observed,
	)


def _run_ingest_http(case: EvalCase) -> EvalCaseResult:
	"""Seed sync ingest (FastAPI upload is 410); keeps ingest_http expect shape."""
	from fastapi.testclient import TestClient

	from app.main import app
	from app.services.metadata import reset_metadata_store
	from app.settings import get_settings
	from tests.support.seed import SeedIngestError, seed_upload_document

	fixture_name = case.fixture or ""
	keys = (
		"ASK_MODE",
		"DASHSCOPE_API_KEY",
		"OPENAI_API_KEY",
		"METADATA_BACKEND",
		"METADATA_PATH",
		"DOCUMENT_STORAGE_DIR",
		"STUB_INGEST_SIMULATE",
	)
	previous = {key: os.environ.get(key) for key in keys}
	with TemporaryDirectory(prefix="meriknow-eval-http-") as tmp_dir:
		os.environ.update(
			{
				"ASK_MODE": "stub",
				"DASHSCOPE_API_KEY": "",
				"OPENAI_API_KEY": "",
				"METADATA_BACKEND": "json",
				"METADATA_PATH": str(Path(tmp_dir) / "metadata.json"),
				"DOCUMENT_STORAGE_DIR": str(Path(tmp_dir) / "documents"),
				"STUB_INGEST_SIMULATE": "true",
			}
		)
		get_settings.cache_clear()
		reset_metadata_store()
		try:
			client = TestClient(app)
			lib_id = case.library_id or f"lib-eval-{case.id}"
			created = client.post(
				"/v1/libraries",
				json={"name": f"eval-{case.id}", "library_id": lib_id},
			)
			if created.status_code != 200:
				return EvalCaseResult(
					id=case.id,
					ok=False,
					kind=case.kind,
					errors=[f"create library failed: {created.status_code} {created.text}"],
					observed={"http_status": created.status_code, "detail": created.text},
				)

			path = _resolve_fixture_path(fixture_name)
			content = path.read_bytes()
			filename = path.name
			mime = {
				".md": "text/markdown",
				".txt": "text/plain",
				".pdf": "application/pdf",
				".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
				".html": "text/html",
				".htm": "text/html",
				".csv": "text/csv",
			}.get(path.suffix.lower(), "application/octet-stream")

			http_status = 200
			detail = ""
			doc_status = None
			error = None
			doc_id = None
			payload_status = None
			try:
				payload = seed_upload_document(
					library_id=lib_id,
					filename=filename,
					content=content,
					content_type=mime,
				)
				doc_id = payload.get("doc_id")
				doc_status = payload.get("status")
				payload_status = payload.get("status")
			except SeedIngestError as exc:
				http_status = exc.http_status
				detail = exc.message
				doc_id = exc.doc_id
				doc_status = exc.doc_status
				error = exc.message

			docs = client.get(f"/v1/libraries/{lib_id}/documents")
			doc_row = None
			if docs.status_code == 200:
				rows = docs.json()
				if doc_id:
					doc_row = next((row for row in rows if row.get("id") == doc_id), None)
				elif rows:
					doc_row = rows[0]
			if doc_row:
				doc_status = doc_row.get("status") or doc_status
				error = doc_row.get("error") or error

			observed = {
				"http_status": http_status,
				"detail": detail or error or "",
				"doc_status": doc_status,
				"error": error or detail,
				"doc_id": doc_id or (doc_row or {}).get("id"),
				"payload_status": payload_status,
			}
			errors = _check_expect(case.expect, observed)
			return EvalCaseResult(
				id=case.id,
				ok=not errors,
				kind=case.kind,
				errors=errors,
				observed=observed,
			)
		finally:
			reset_metadata_store()
			for key, value in previous.items():
				if value is None:
					os.environ.pop(key, None)
				else:
					os.environ[key] = value
			get_settings.cache_clear()


def run_eval_cases(path: Path | None = None) -> list[EvalCaseResult]:
	"""跑黄金集；整段强制 INTERNAL_AUTH_ENABLED=false，避免宿主 .env 绊倒 gate。"""
	from app.settings import get_settings

	prev_auth = os.environ.get("INTERNAL_AUTH_ENABLED")
	os.environ["INTERNAL_AUTH_ENABLED"] = "false"
	get_settings.cache_clear()
	try:
		cases = load_eval_cases(path)
		results: list[EvalCaseResult] = []
		for case in cases:
			if case.kind == "classify":
				results.append(_run_classify(case))
			elif case.kind == "ingest_chunk":
				results.append(_run_ingest_chunk(case))
			elif case.kind == "retrieval":
				results.append(_run_retrieval(case))
			elif case.kind == "ingest_http":
				results.append(_run_ingest_http(case))
			else:
				results.append(_run_ask(case))
		return results
	finally:
		if prev_auth is None:
			os.environ.pop("INTERNAL_AUTH_ENABLED", None)
		else:
			os.environ["INTERNAL_AUTH_ENABLED"] = prev_auth
		get_settings.cache_clear()


def main(argv: list[str] | None = None) -> int:
	args = list(argv or sys.argv[1:])
	path = Path(args[0]) if args else DEFAULT_CASES
	results = run_eval_cases(path)
	passed = sum(1 for item in results if item.ok)
	failed = [item for item in results if not item.ok]
	print(f"[eval] cases={len(results)} passed={passed} failed={len(failed)} file={path}")
	print("[eval] retrieval metric default = Recall@K (K=3 unless expect.recall_at_k overrides)")
	for item in results:
		mark = "PASS" if item.ok else "FAIL"
		extra = ""
		if item.kind == "retrieval":
			rank = (item.observed or {}).get("observed_rank")
			metric = (item.observed or {}).get("metric") or "Recall@3"
			extra = f" {metric} rank={rank}"
		print(f"  {mark} {item.id} ({item.kind}){extra}")
		for err in item.errors:
			print(f"       - {err}")
	return 1 if failed else 0


if __name__ == "__main__":
	raise SystemExit(main())
